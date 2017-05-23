#!/usr/bin/env python
import logging
import re
import tempfile
from subprocess import call

import pdfkit
from docx import Document

from classes.exceptions import *
from classes.utils import IPFSDocument
from config import *


class PdfFactory(object):
    def factory(file_type, hash_key, replace_tags=None):

        if file_type not in DOCUMENT_RENDERING_OPTIONS:
            raise UnSupportedFileException("Unsupported extension %s" % file_type)

        document_rendered_options = DOCUMENT_RENDERING_OPTIONS[file_type]

        if file_type == "html":
            return HtmlDocument(hash_key=hash_key, replace_tags=replace_tags,
                                document_rendered_options=document_rendered_options)
        elif file_type == "word":
            return WordDocument(hash_key=hash_key, replace_tags=replace_tags,
                                document_rendered_options=document_rendered_options)

    factory = staticmethod(factory)


# Decorator : If pdf file exist no need to execute the function.
def check_pdf_file_exists(func):
    def checker(*args):
        file_path = args[0].pdf_file_path
        if os.path.exists(file_path):
            logging.info("The pdf file exists before in path %s" % file_path)
            return file_path

        return func(*args)

    return checker


class WordDocument(IPFSDocument):
    def __init__(self, hash_key, replace_tags=None, document_rendered_options=None):
        super(WordDocument, self).__init__(hash_key=hash_key, replace_tags=replace_tags, extension="word")
        self.temp_word_file_with_tags_replaced = None

    def _paragraph_replace(self, tags_dic):
        pattern = re.compile('|'.join(tags_dic.keys()))
        for paragraph in self.doc.paragraphs:
            if paragraph.text:
                paragraph.text = pattern.sub(lambda m: tags_dic[m.group(0)], paragraph.text)

    def _replace_tags(self, word_file_path):
        if os.path.exists(word_file_path):
            logging.debug('opening source word document: %s' % word_file_path)

            self.doc = Document(word_file_path)
            if self.replace_tags:
                logging.debug('start replacing tags: %s' % word_file_path)
                self._paragraph_replace(self.replace_tags)

            temp = tempfile.NamedTemporaryFile(prefix='document_word_', dir=IPFS_CACHE_DIR, delete=False)
            with temp:
                self.temp_word_file_with_tags_replaced = temp.name  # Temporary file name
            logging.debug('temporary file with tags replaces in %s' % self.temp_word_file_with_tags_replaced)

            # Replace the downloaded temp file with the new document with tags.
            self.doc.save(self.temp_word_file_with_tags_replaced)

            logging.debug('file has been saved in: %s/%s' % (CONVERTED_DIR, self.encoded_hash))

            return self.temp_word_file_with_tags_replaced
        else:
            raise NotFoundException("Docx file %s.%s not found" % (self.encoded_hash, self.extension))

    def _word_pdf(self, word_file_path):
        logging.info("start pdf:%s" % word_file_path)

        if os.path.exists(word_file_path):
            script = '%s/doc2pdf.sh "%s" "%s"' % (CURRENT_DIRECTORY, word_file_path, CONVERTED_DIR)
            logging.debug("Execute script %s" % script)

            # execute bash script doc2pdf using `soffice` command
            result = call(script, shell=True)
            if result == 0:
                old_file_name = os.path.basename(os.path.normpath(self.temp_word_file_with_tags_replaced))
                expected_pdf_file = "%s/%s.pdf" % (CONVERTED_DIR, old_file_name)
                if os.path.exists(expected_pdf_file):
                    logging.info("PDF file saved successfully")
                    return expected_pdf_file
                else:
                    err_message = "Bash script error, No PDF file %s exists" % expected_pdf_file
                    raise BashScriptException(err_message)
            else:
                err_message = "Bash script error in saving PDF file with hash %s" % self.encoded_hash
                raise BashScriptException(err_message)
        else:
            raise NotFoundException("%s File is not found" % (self.encoded_hash))

    @check_pdf_file_exists
    def generate(self):
        word_tags_replaced_file_path = None
        try:
            # download ipfs document if not exists in cache
            word_ipfs_file_path = self.download_pinned_ipfs_document_into_cache()

            # Replaced tags if exists, save replaced document in temp file
            word_tags_replaced_file_path = self._replace_tags(word_ipfs_file_path)

            # convert the replaced document file into pdf, then save it into converted folder
            temporary_pdf_file_path = self._word_pdf(word_tags_replaced_file_path)

            # rename the pdf file in converted folder into encoded hash to use in in cache
            os.rename(temporary_pdf_file_path, self.pdf_file_path)

        finally:
            # Remove the temp file that replaced the tags if exist
            if word_tags_replaced_file_path and os.path.exists(word_tags_replaced_file_path):
                os.remove(word_tags_replaced_file_path)

        return self.pdf_file_path


class HtmlDocument(IPFSDocument):
    def __init__(self, hash_key, replace_tags=None, document_rendered_options=None):
        super(HtmlDocument, self).__init__(hash_key=hash_key, replace_tags=replace_tags, extension="html")
        self.document_rendered_options = document_rendered_options

    def _replace_tags(self, html_file_path):
        # open html file from temp folder
        with open(html_file_path, "r") as file:
            data = file.read()

        # check if there's tags needs to replace
        if self.replace_tags:
            pattern = re.compile('|'.join(self.replace_tags.keys()))
            data = pattern.sub(lambda m: self.replace_tags[m.group(0)], str(data))

        return data

    def _html_pdf(self, string_html, pdf_folder):
        pdfkit.from_string(string_html, pdf_folder, options=self.document_rendered_options)
        if not os.path.exists(pdf_folder):
            raise FileNotFoundError("Fail to create %s file" % pdf_folder)

    @check_pdf_file_exists
    def generate(self):
        # download ipfs document if not exists before in the temp folder
        html_ipfs_file_path = self.download_pinned_ipfs_document_into_cache()

        # Replaced tags if exists, return the results into data variable
        data = self._replace_tags(html_ipfs_file_path)

        # convert the result in variable data into pdf file, save it into converted folder
        self._html_pdf(str(data), self.pdf_file_path)

        return self.pdf_file_path
